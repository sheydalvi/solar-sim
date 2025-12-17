import time
start_time = time.time()

import pandas as pd
from core.SolarSimScripts import *

from datetime import date
import csv
import streamlit as st


SM = 0
NU = 0
TI = 0


def cell_format_data(table):
    """
    Formats the second column in a table with given text, no space after
    """
    # Loop through each row in the table
    for row in table.rows:
        cell = row.cells[1]

        # Add paragraph and apply alignment
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(0)

            # Apply formatting to all runs
            for run in paragraph.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(11)
    
                # Fix font name for Word compatibility
                r = run._element
                r.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')


def cell_format_info(table, row, col, text):
    """
    Formats a specific cell in a table with given text, no space after
    """
    cell = table.cell(row, col)
    paragraph = cell.paragraphs[0]
    paragraph.clear() 
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(11)



st.title("SciSun Test Analysis Dashbord")

col1, col2, col3 = st.columns(3)
with col1:
    PN = st.text_input(
    "Project Number:",
    "",
    key="pn",
)
with col2:
    SN = st.text_input(
    "SciSun Serial Number:",
    "",
    key="sn",
)
with col3:
    PS = st.text_input(
    "Power Supply Serial Number:",
    "",
    key="ps",
)
with col1:
    user = st.text_input(
    "End User:",
    "",
    key="user",
)
with col2:
    Country = st.text_input(
    "End User Country:",
    "",
    key="country",
)
with col3:
    tester = st.text_input(
    "Test Performed by:",
    "SA",
    key="tester",
)
    

# with col2:
#     Scisun = st.selectbox("SciSun Model:", models)

    
models= ['[160-9108-C01] Scisun-150-1.5G-TS-MS-NA',
'[160-9108-C02] Scisun-150-1.5G-TS-MS-AA',
'[160-9108-C03] Scisun-150-1.5G-TS-MS-MA',
'[160-9108-C04] Scisun-150-1.5G-TS-AS-NA',
'[160-9108-C05] Scisun-150-1.5G-TS-AS-AA',
'[160-9108-C06] Scisun-150-1.5G-TS-AS-MA',
'[160-9108-C07] Scisun-150-1.5G-MC-MS-NA',
'[160-9108-C08] Scisun-150-1.5G-MC-MS-AA',
'[160-9108-C09] Scisun-150-1.5G-MC-MS-MA',
'[160-9108-C13] Scisun-300-1.5G-TS-MS-NA',
'[160-9108-C14] Scisun-300-1.5G-TS-MS-AA',
'[160-9108-C15] Scisun-300-1.5G-TS-MS-MA',
'[160-9108-C16] Scisun-300-1.5G-TS-AS-NA',
'[160-9108-C17] Scisun-300-1.5G-TS-AS-AA',
'[160-9108-C18] Scisun-300-1.5G-TS-AS-MA',
'[160-9108-C19] Scisun-300-1.5G-MC-MS-NA',
'[160-9108-C20] Scisun-300-1.5G-MC-MS-AA',
'[160-9108-C21] Scisun-300-1.5G-MC-MS-MA',
'(160-9109-C01) SCISUN-300-0.00-TS-MS-NA-AAA',
'(160-9109-C02) SCISUN-300-0.00-TS-MS-MA-AAA',
'(160-9109-C03) SCISUN-300-0.00-TS-MS-AA-AAA',
'(160-9109-C04) SCISUN-300-0.00-TS-AS-NA-AAA',
'(160-9109-C05) SCISUN-300-0.00-TS-AS-MA-AAA',
'(160-9109-C06) SCISUN-300-0.00-TS-AS-AA-AAA',
'(160-9109-C07) SCISUN-300-0.00-TS-MS-NA-BAA',
'(160-9109-C08) SCISUN-300-0.00-TS-MS-MA-BAA',
'(160-9109-C09) SCISUN-300-0.00-TS-MS-AA-BAA',
'(160-9109-C10) SCISUN-300-0.00-TS-AS-NA-BAA',
'(160-9109-C11) SCISUN-300-0.00-TS-AS-MA-BAA',
'(160-9109-C12) SCISUN-300-0.00-TS-AS-AA-BAA',
'[160-9109-C13] SCISUN-300-0.00-MC-MS-NA-AAA',
'[160-9109-C14] SCISUN-300-0.00-MC-MS-MA-AAA',
'[160-9109-C15] SCISUN-300-0.00-MC-MS-AA-AAA',
'[160-9109-C16] SCISUN-300-0.00-MC-MS-NA-BAA',
'[160-9109-C17] SCISUN-300-0.00-MC-MS-MA-BAA',
'[160-9109-C18] SCISUN-300-0.00-MC-MS-AA-BAA']
Scisun = st.selectbox("SciSun Model:", models)
psmodels= ['601-300','601-150']
Psmodel = st.selectbox("Power Supply Model:", psmodels)




st.header("Measurement Equipment")
df = pd.DataFrame({
    "Measurement Equipment": ["Solar Reference Cell", "Source Meter", "Spectroradiometer Model", "Temperature Sensor"],
    "Model": ["SC-LT-Q", "Keithley 2400", "9055-S-SOLARX/SD/P", "Amprobe TH-1"],
    "Serial No.": ["", "", "", ""],
    "Calibration Date": ["", "", "", ""],
})

filled_out = st.data_editor(df, num_rows="dynamic")


st.header("Solar Simulator Specifications")


spectab = pd.DataFrame({
    "spec": ["filter", "1 Sun", "2 Sun", "Max", "Min", "field of view", "Lamp age", "warmup", "T", "date", "target area", "A", "B"],
    "value": ["", "", "", "", "", "8.16", "", "15", "", date.today().strftime("%Y-%m-%d"), "50x50", "380", "4"],
    "notes": ["AM Filter Part Number: XXX-XXXX\nBatch: XXXXX",
              "Attenuator Setting: XX\nMesh Filter: XX",
              "Attenuator Setting: XX\nMesh Filter: XX",
              "Power Supply Setting: XX.X%\nAttenuator Setting: XX\nMesh Filter: XX",
              "Power Supply Setting: XX.X%\nAttenuator Setting: XX\nMesh Filter: XX",
              "","", "","","","","",""],
})

columns = ['Model', 'Serial No.', 'Calibration Date']
filled_out_specs = st.data_editor(spectab, num_rows="dynamic")



# --- NU ---
st.header("Non-Uniformity")

uploaded_NU = st.file_uploader("Upload .sudat file", type=['sudat'])

if uploaded_NU is not None:
    NU = 1
    from core.SciImports import sudatImport
    mes_date_nu = st.date_input("Enter the NU masurement date")
    result_NU = sudatImport(uploaded_NU)

    decoded_data_NU = uploaded_NU.getvalue().decode('utf-8')
    comma_delimited_data_NU = decoded_data_NU.replace('\t', ',')
    csv_bytes_NU = comma_delimited_data_NU.encode('utf-8')

    st.success("File successfully parsed.")

    # Optional: Call your analysis
    NU_report = NUScript(result_NU)
    st.write("Results:", NU_report)

else:
    st.error("NU file not uploaded")


# --- TI ---
st.header("Temporal Instability")

uploaded_TI = st.file_uploader("Upload .ivdat file", type=['ivdat'], accept_multiple_files=True)

if uploaded_TI:
    TI = 1
    from core.SciImports import ivdatImports
    result_TI = ivdatImports(uploaded_TI)
    mes_date_ti = st.date_input("Enter the TI masurement date")


    # Run analysis
    TI_report = TIScript(result_TI)
    worst_filename = TI_report['Filename']

    # Find the corresponding original file
    if isinstance(uploaded_TI, list):  # multiple file upload
        for file in uploaded_TI:
            if file.name == worst_filename:
                worst_file_obj = file
                break
        else:
            st.error("Original file for worst dataset not found.")
            st.stop()
    else:  # single file upload
        worst_file_obj = uploaded_TI

    # Convert tab-delimited to comma-delimited
    decoded_data_TI = worst_file_obj.getvalue().decode('utf-8')
    comma_delimited_data_TI = decoded_data_TI.replace('\t', ',')
    csv_bytes_TI = comma_delimited_data_TI.encode('utf-8')

    # Offer download
    # st.download_button(
    #     label="Download Worst TI CSV",
    #     data=csv_bytes_TI,
    #     file_name=f"TemporalInstabilityData_ProjectNo{PN}SN{SN}.csv",
    #     mime="text/csv"
    # )

    st.success("Files successfully parsed.")

    # Display parsed info (optional)
    # for key, value in result_TI.items():
    #     st.subheader(value['filename'])
    #     st.write(value)

    st.write("Results:", TI_report)
else:
    st.error("TI file not uploaded")


# --- SM ---
st.header("Spectrometry Analysis")

standards = {
    "AM1.5G Hemispherical ASTM E927-19": '2',
    "AM1.5D Direct Normal ASTM E927-19": '1',
    "AM0 Extra-Terrestrial ASTM E927-19": '3',
    "AM1.5G IEC 60904-9 Ed.3 Table 1": '4',
    "AM1.5G IEC 60904-9 Ed. 3 Table 2": '5',
    "AM1.5G Hemispherical ASTM E927-19, Limited Range [700 - 1100 nm]": '6',
}


col11, col22 = st.columns(2) 
with col11:
    user_am = st.selectbox("The Standard:", list(standards.keys()))
    AMType = standards[user_am]  


label = st.text_input(
    "Label the data for the plot:",
    "SciSun300 SN0001234",
    key="label",
)


uploaded_full = st.file_uploader("Upload .sidat file full", type=['sidat'], key='full')


if uploaded_full:
    SM= 1
    from core.SciImports import sidatImport
    mes_date_sm = st.date_input("Enter the SM masurement date")
    result= sidatImport(uploaded_full)
    st.success("File successfully parsed.")



    SM_report = SMScript2(AMType, result, label)
    st.write("Results:", SM_report)

    df = SM_report['df']
    csv_data_SM = df.to_csv(index=False).encode('utf-8')
else:
    st.error("SM file not uploaded")





if SM and TI and NU:
    import matplotlib.pyplot as plt
    from docx import Document
    from docx.shared import Inches
    from docx.shared import RGBColor
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn

    doc =  Document(r"\\10.0.0.11\TestingLibrary\TESTING REPORTS\TEMPLATES\160- SMALL SOLARS\TR-SCISUN-01 E927-19_pythontemplate.docx")

    # access a specific table by index
    projrect = doc.tables[0]  
    model = doc.tables[1]  
    performer = doc.tables[2]  
    equip = doc.tables[3]  
    specs = doc.tables[4]
    nonuni = doc.tables[5]      # NU
    spectral = doc.tables[6]    # SM  
    tempinst = doc.tables[7]    # TI  


    cell_format_info(projrect, 0, 1, PN)
    cell_format_info(projrect, 1, 1, user)
    cell_format_info(projrect, 2, 1, Country)


    cell_format_info(model, 0, 1, Scisun)
    cell_format_info(model, 1, 1, SN)
    cell_format_info(model, 2, 1, Psmodel)
    cell_format_info(model, 3, 1, PS)

    cell_format_info(performer, 0, 1, tester)
    cell_format_info(performer, 1, 1, date.today().strftime("%Y-%m-%d"))





    for row in range(4): 
        for col_index, column_name in enumerate(columns, start=1):
            cell_format_info(equip, row + 1, col_index, filled_out.loc[row, column_name])


    cell_format_info(specs, 1, 1, filled_out_specs.loc[0, 'value'])
    cell_format_info(specs, 2, 1, filled_out_specs.loc[1, 'value'])
    cell_format_info(specs, 3, 1, filled_out_specs.loc[2, 'value'])
    cell_format_info(specs, 4, 1, filled_out_specs.loc[3, 'value'])
    cell_format_info(specs, 5, 1, filled_out_specs.loc[4, 'value'])
    cell_format_info(specs, 6, 1, filled_out_specs.loc[5, 'value'])
    cell_format_info(specs, 7, 1, filled_out_specs.loc[6, 'value'])
    cell_format_info(specs, 8, 1, filled_out_specs.loc[7, 'value'])
    cell_format_info(specs, 9, 1, filled_out_specs.loc[8, 'value'])
    cell_format_info(specs, 10, 1, filled_out_specs.loc[9, 'value'])



    cell_format_info(specs, 1, 3, filled_out_specs.loc[0, 'notes'])
    cell_format_info(specs, 2, 3, filled_out_specs.loc[1, 'notes'])
    cell_format_info(specs, 3, 3, filled_out_specs.loc[2, 'notes'])
    cell_format_info(specs, 4, 3, filled_out_specs.loc[3, 'notes'])
    cell_format_info(specs, 5, 3, filled_out_specs.loc[4, 'notes'])
    cell_format_info(specs, 6, 3, filled_out_specs.loc[5, 'notes'])
    cell_format_info(specs, 7, 3, filled_out_specs.loc[6, 'notes'])
    cell_format_info(specs, 8, 3, filled_out_specs.loc[7, 'notes'])
    cell_format_info(specs, 9, 3, filled_out_specs.loc[8, 'notes'])


    # TABLE
    nonuni.cell(1, 1).text = str(mes_date_nu)        # [YYYY-MM-DD]
    nonuni.cell(2, 1).text = '0.3' # NU_report['Measurement Grid Size [cm]']            # Detector Area:
    nonuni.cell(3, 1).text = str(NU_report['Number of Measurement Points'])            # Number of Measurement Points:
    nonuni.cell(4, 1).text = f"{float(NU_report['Ideal Detector Area [cm^2]']):.3f}"             # Measurement Point Area:
    nonuni.cell(5, 1).text = f"{float(NU_report['Maximum Irradiance [Suns]']):.3f}"            # Maximum Irradiance:
    nonuni.cell(6, 1).text = f"{float(NU_report['Minimum Irradiance [Suns]']):.3f}"            # Minimum Irradiance:
    nonuni.cell(7, 1).text = f"{float(NU_report['Standard Deviation [Suns]']):.3f}"            # Sample Standard Deviation of Spatial Non-Uniformity:
    nonuni.cell(8, 1).text = f"{float(NU_report['Spatial Non-Uniformity of Irradiance [%]']):.3f}"           # Spatial Non-Uniformity of Irradiance:
    # Classification:
    if NU_report['Spatial Non-Uniformity of Irradiance [%]'] <= 2.0:
        nonuni.cell(9, 1).text = "A"         # Classification:
    elif NU_report['Spatial Non-Uniformity of Irradiance [%]'] <= 5.0:
        nonuni.cell(9, 1).text = "B"         
    else:
        nonuni.cell(9, 1).text = "C"         
    cell_format_data(nonuni)

    # PLOT
    # search for the placeholder and replace with the image
    for i, para in enumerate(doc.paragraphs):
        if "{{INSERT_PLOT_NU}}" in para.text:
            para.text = para.text.replace("{{INSERT_PLOT_NU}}", "")
            
            # insert image right after the current paragraph
            run = para.insert_paragraph_before().add_run()
            run.add_picture("output/NU.png", width=Inches(7))
            break


    # TABLE
    tempinst.cell(1, 1).text = str(mes_date_ti)                                  # [YYYY-MM-DD]
    # tempinst.cell(2, 1).text = str(TI_report['Detector Area'])                    # Detector Area:
    tempinst.cell(3, 1).text = str(TI_report['Time Between Data Points [s]'])                     # Number of Measurement Points:
    tempinst.cell(4, 1).text = str(TI_report['Number of Power Line Cycles'])             # Measurement Point Area:
    tempinst.cell(5, 1).text = str(TI_report['Total Measurement Points'])          # Maximum Irradiance:
    tempinst.cell(6, 1).text = f"{float(TI_report['Minimum Irradiance [Suns]']):.3f}"            # Maximum Irradiance:
    tempinst.cell(7, 1).text = f"{float(TI_report['Maximum Irradiance [Suns]']):.3f}"            # Minimum Irradiance:
    tempinst.cell(8, 1).text = f"{float(TI_report['Temporal Instability [%]']):.3f}"           # Spatial Non-Uniformity of Irradiance:
    # Classification:
    if TI_report['Temporal Instability [%]'] <= 4:
        tempinst.cell(9, 1).text = "A"         # Classification:
    elif TI_report['Temporal Instability [%]'] <= 10:
        tempinst.cell(9, 1).text = "B"         
    else:
        tempinst.cell(9, 1).text = "C"         
    cell_format_data(tempinst)

    # PLOT
    # search for the placeholder and replace with the image
    for i, para in enumerate(doc.paragraphs):
        if "{{INSERT_PLOT_TI}}" in para.text:
            para.text = para.text.replace("{{INSERT_PLOT_TI}}", "")
            
            # insert image right after the current paragraph
            run = para.insert_paragraph_before().add_run()
            run.add_picture("output/TI.png", width=Inches(7))
            break

    # TABLE
    spectral.cell(1, 1).text = str(mes_date_sm)
    spectral.cell(2, 1).text = '20' # str(SM_report['Scan Time'])                
    spectral.cell(3, 1).text = '1' # str(SM_report['number of spectra avg'])                 
    spectral.cell(4, 1).text = SM_report['Classification']                
    cell_format_data(spectral)

    # PLOT
    # search for the placeholder and replace with the image
    for i, para in enumerate(doc.paragraphs):
        if "{{INSERT_PLOT_SM}}" in para.text:
            para.text = para.text.replace("{{INSERT_PLOT_SM}}", "")
            
            # insert image right after the current paragraph
            run = para.insert_paragraph_before().add_run()
            run.add_picture("output/SM.png", width=Inches(7))
            break

    col111, col222, col333, col444 = st.columns(4)
    with col111:
        if st.button("Generate Report"):
            import io
            # save to in-memory file
            doc_io = io.BytesIO()
            doc.save(doc_io)
            doc_io.seek(0)

            st.download_button(
                label="Download Report as .docx",
                data=doc_io,
                file_name=f"TR-SCISUN-01 E927-19 [ProjectNo{PN} SN{SN}].docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

    with col222:
        # generate and download and  your files
        st.download_button("Download NU file", data=csv_bytes_NU, file_name=f"SpatialUniformityData_ProjectNo{PN}SN{SN}.csv")
    with col333:
        st.download_button("Download TI file", data=csv_bytes_TI, file_name=f"TemporalInstabilityData_ProjectNo{PN}SN{SN}.csv")
    with col444:
        st.download_button("Download SM file", data=csv_data_SM, file_name=f"SpectralMatchData_ProjectNo{PN}SN{SN}.csv")
